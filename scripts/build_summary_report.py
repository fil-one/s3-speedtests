#!/usr/bin/env python3
"""Build a Google Docs-friendly DOCX or PDF summary from benchmark JSONL output."""

from __future__ import annotations

import argparse
import configparser
import datetime as dt
import html
import json
import os
import re
from pathlib import Path
import subprocess
import sys
from typing import Any
from urllib.parse import urlparse

PROVIDERS = {
    "f1": ("Fil.one", "eu-west-1", "Albi, France"),
    "fil.one": ("Fil.one", "eu-west-1", "Albi, France"),
    "filone": ("Fil.one", "eu-west-1", "Albi, France"),
    "aws": ("AWS", "eu-south-2", "Aragon, Spain"),
    "aws-test": ("AWS", "eu-south-2", "Aragon, Spain"),
    "wasabi": ("Wasabi", "eu-west-2", "Paris, France"),
    "backblaze": ("Backblaze", "eu-central-003", "Amsterdam, Netherlands"),
}

REGION_LOCATIONS = {
    "eu-west-1": "Albi, France",
    "eu-south-2": "Aragon, Spain",
    "eu-west-3": "Paris, France",
    "eu-west-2": "Paris, France",
    "eu-central-003": "Amsterdam, Netherlands",
    "us-west-2": "Oregon, USA",
}

PROVIDER_NAMES = {
    "f1": "Fil.one",
    "fil.one": "Fil.one",
    "filone": "Fil.one",
    "aws": "AWS",
    "aws-test": "AWS",
    "wasabi": "Wasabi",
    "backblaze": "Backblaze",
}

FILONE_BRAND_FILL = "C9DAF8"
BAR_SEGMENTS = 16
PROVIDER_BAR_FILLS = {
    "aws": "FF9900",
    "aws-test": "FF9900",
    "backblaze": "E21E29",
    "f1": FILONE_BRAND_FILL,
    "fil.one": FILONE_BRAND_FILL,
    "filone": FILONE_BRAND_FILL,
    "wasabi": "17D24F",
}
DEFAULT_BAR_FILL = "DADCE0"


def import_docx() -> None:
    global Document
    global Inches
    global OxmlElement
    global Pt
    global RGBColor
    global WD_ALIGN_PARAGRAPH
    global WD_CELL_VERTICAL_ALIGNMENT
    global WD_SECTION
    global WD_TABLE_ALIGNMENT
    global qn

    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise SystemExit("python-docx is required. Run scripts/setup_vm.sh or install python3-docx.") from exc


def import_reportlab() -> None:
    global colors
    global getSampleStyleSheet
    global inch
    global landscape
    global letter
    global PageBreak
    global Paragraph
    global ParagraphStyle
    global PdfTable
    global Preformatted
    global SimpleDocTemplate
    global Spacer
    global TableStyle

    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table as PdfTable, TableStyle
    except ModuleNotFoundError as exc:
        raise SystemExit("python3-reportlab is required for PDF output. Run scripts/setup_vm.sh or install python3-reportlab.") from exc


def provider_key(value: str) -> str:
    return value.strip().lower()


def bool_value(value: str | None, default: bool = False) -> bool:
    if value is None or value == "":
        return default
    return value.strip().lower() in {"1", "yes", "true", "on", "enable", "enabled"}


def load_jsonl(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def latest(data_dir: Path, pattern: str, fallback: str) -> Path:
    matches = sorted(data_dir.glob(pattern), key=lambda p: p.stat().st_mtime, reverse=True)
    return matches[0] if matches else data_dir / fallback


def first_config_value(parser: configparser.ConfigParser, section: str, keys: list[str]) -> str:
    for key in keys:
        value = parser.get(section, key, fallback="").strip()
        if value:
            return value
    return ""


def endpoint_host(endpoint_url: str) -> str:
    endpoint_url = endpoint_url.strip()
    if not endpoint_url:
        return ""
    parsed = urlparse(endpoint_url if "://" in endpoint_url else f"https://{endpoint_url}")
    return (parsed.netloc or parsed.path).split("/")[0]


def derived_endpoint(provider: str, region: str, endpoint_url: str) -> str:
    host = endpoint_host(endpoint_url)
    if host:
        return host
    provider = provider_key(provider)
    if provider in {"aws", "aws-test"} and region:
        return f"s3.{region}.amazonaws.com"
    if provider == "wasabi" and region:
        return f"s3.{region}.wasabisys.com"
    if provider == "backblaze" and region:
        return f"s3.{region}.backblazeb2.com"
    if provider in {"f1", "fil.one", "filone"} and region:
        return f"{region}.s3.fil.one"
    return ""


def load_provider_labels(targets_path: Path) -> dict[str, dict[str, Any]]:
    labels: dict[str, dict[str, Any]] = {}
    if not targets_path.exists():
        return labels

    parser = configparser.ConfigParser(interpolation=None)
    parser.read(targets_path)
    for section in parser.sections():
        provider = provider_key(parser.get(section, "provider", fallback=section))
        bucket = parser.get(section, "bucket", fallback="").strip()
        region = parser.get(section, "region", fallback="").strip()
        endpoint_url = parser.get(section, "endpoint_url", fallback="").strip()
        enabled = bool_value(parser.get(section, "enabled", fallback="false"))
        endpoint = derived_endpoint(provider, region, endpoint_url)
        name = first_config_value(parser, section, ["display_name", "provider_name", "name"]) or PROVIDER_NAMES.get(provider, section)
        location = first_config_value(parser, section, ["location", "region_location", "city_country"]) or REGION_LOCATIONS.get(region, "")
        label = {
            "name": name,
            "region": region,
            "location": location,
            "bucket": bucket,
            "section": section,
            "endpoint_url": endpoint_url,
            "endpoint": endpoint,
            "enabled": enabled,
        }
        labels[f"section:{section}"] = label
        labels[f"provider:{provider}"] = label
        if bucket:
            labels[f"bucket:{bucket}"] = label
        if provider and bucket:
            labels[f"provider_bucket:{provider}:{bucket}"] = label
    return labels


def enabled_traceroute_endpoints(provider_labels: dict[str, dict[str, Any]]) -> set[str]:
    endpoints = {
        str(label.get("endpoint") or "").lower()
        for label in provider_labels.values()
        if label.get("enabled") and label.get("endpoint")
    }
    return {endpoint for endpoint in endpoints if endpoint}


def filter_traceroute_records(records: list[dict[str, Any]], endpoints: set[str]) -> list[dict[str, Any]]:
    rows = [r for r in records if r.get("test_type") == "tcp_traceroute_443"]
    if not endpoints:
        return rows
    return [r for r in rows if str(r.get("endpoint") or "").lower() in endpoints]


def command_text(cmd: list[str]) -> str:
    try:
        return subprocess.check_output(cmd, stderr=subprocess.DEVNULL, text=True).strip()
    except (FileNotFoundError, subprocess.CalledProcessError):
        return ""


def detect_hostname() -> str:
    return command_text(["hostname"]) or os.environ.get("HOSTNAME", "") or "unknown"


def detect_vcpus() -> str:
    value = command_text(["nproc"])
    if not value:
        value = command_text(["sh", "-c", "getconf _NPROCESSORS_ONLN"])
    return f"{value} vCPU" if value else "unknown"


def detect_ram() -> str:
    value = command_text(["sh", "-c", "free -m | awk '/^Mem:/ {print $2}'"])
    if not value:
        value = command_text(["sh", "-c", "awk '/MemTotal:/ {printf \"%.0f\", $2 / 1024}' /proc/meminfo"])
    if not value:
        return "unknown"
    try:
        gib = float(value) / 1024
    except ValueError:
        return "unknown"
    if gib >= 10:
        return f"{gib:.0f} GB RAM"
    return f"{gib:.1f} GB RAM"


def prompt_value(label: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    value = input(f"{label}{suffix}: ").strip()
    return value or default


def resolve_source_context(args: argparse.Namespace) -> argparse.Namespace:
    args.node_hostname = args.node_hostname or detect_hostname()
    args.node_compute = args.node_compute or detect_vcpus()
    args.node_memory = args.node_memory or detect_ram()

    provider_default = args.source_provider or os.environ.get("SOURCE_NODE_PROVIDER", "")
    location_default = args.source_location or os.environ.get("SOURCE_NODE_LOCATION", "")

    if not args.no_prompt and sys.stdin.isatty():
        args.source_provider = prompt_value("Source node provider/name", provider_default)
        args.source_location = prompt_value("Source node location", location_default)
    else:
        args.source_provider = provider_default
        args.source_location = location_default

    args.source_provider = args.source_provider or args.node_hostname or "unknown"
    args.source_location = args.source_location or "unknown"
    return args


def fmt_size(mib: float) -> str:
    if mib >= 1024:
        gib = mib / 1024
        return f"{int(round(gib))} GiB" if abs(gib - round(gib)) < 0.001 else f"{gib:.1f} GiB"
    return f"{mib:g} MiB"


def fmt_mbps(value: float | None) -> str:
    if value is None:
        return "n/a"
    return f"{value:,.0f} Mbps" if value >= 100 else f"{value:,.2f} Mbps"


def fmt_seconds(value: Any) -> str:
    if value is None or value == "":
        return "n/a"
    try:
        seconds = float(value)
    except (TypeError, ValueError):
        return str(value)
    if seconds >= 3600:
        return f"{seconds / 3600:.2f} h"
    if seconds >= 60:
        return f"{seconds / 60:.2f} min"
    return f"{seconds:.3f} s"


def fmt_ms(value: Any) -> str:
    if value is None or value == "":
        return "n/a"
    try:
        return f"{float(value):.2f} ms"
    except (TypeError, ValueError):
        return str(value)


def traceroute_total_ms(record: dict[str, Any]) -> float | None:
    for key in ("total_ms", "final_hop_avg_ms", "last_responding_hop_avg_ms"):
        value = record.get(key)
        if value is not None and value != "":
            try:
                return float(value)
            except (TypeError, ValueError):
                pass

    trace_output = str(record.get("trace_output") or "")
    for line in reversed(trace_output.splitlines()):
        samples = [float(match.group(1)) for match in re.finditer(r"([0-9]+(?:\.[0-9]+)?)\s*ms\b", line)]
        if samples:
            return sum(samples) / len(samples)
    return None


def fmt_trace_ips(value: Any) -> str:
    if not value:
        return "n/a"
    if isinstance(value, list):
        return "\n".join(str(item) for item in value[:3]) + ("\n..." if len(value) > 3 else "")
    return str(value)


def label_from_row(row: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> dict[str, Any]:
    provider = provider_key(str(row.get("provider", "")))
    bucket = str(row.get("bucket", "")).strip()
    section = str(row.get("target_section", "")).strip()
    for key in (
        f"section:{section}" if section else "",
        f"provider_bucket:{provider}:{bucket}" if provider and bucket else "",
        f"bucket:{bucket}" if bucket else "",
        f"provider:{provider}" if provider else "",
    ):
        if key and key in provider_labels:
            return provider_labels[key]

    default_name, default_region, default_location = PROVIDERS.get(provider, (row.get("provider") or "Provider", "", ""))
    region = str(row.get("region") or default_region or "").strip()
    return {
        "name": PROVIDER_NAMES.get(provider, default_name),
        "region": region,
        "location": REGION_LOCATIONS.get(region, default_location if region == default_region else ""),
        "bucket": bucket,
        "section": section,
        "endpoint_url": row.get("endpoint_url", ""),
    }


def provider_label(row: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> str:
    label = label_from_row(row, provider_labels)
    region = str(label.get("region") or "").strip()
    location = str(label.get("location") or "").strip()
    if region and location:
        return f"{label['name']}\n({region} | {location})"
    if region:
        return f"{label['name']}\n({region})"
    return str(label["name"])


def provider_chart_label(row: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> str:
    return provider_label(row, provider_labels).replace("\n", " ")


def provider_bar_fill(row: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> str:
    provider = provider_key(str(row.get("provider", "")))
    if is_filone_row(row, provider_labels):
        return FILONE_BRAND_FILL
    return PROVIDER_BAR_FILLS.get(provider, DEFAULT_BAR_FILL)


def compact_provider(provider: str) -> str:
    key = provider_key(provider)
    return PROVIDER_NAMES.get(key, PROVIDERS.get(key, (provider, "", ""))[0])


def is_filone_row(row: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> bool:
    provider = provider_key(str(row.get("provider", "")))
    if provider in {"f1", "fil.one", "filone"}:
        return True
    return str(label_from_row(row, provider_labels).get("name") or "").strip().lower() == "fil.one"


def summary_total_elapsed(record: dict[str, Any]) -> float | None:
    value = record.get("total_elapsed_seconds")
    if value is not None:
        return value
    avg_elapsed = record.get("avg_elapsed_seconds")
    successes = record.get("success_count")
    if avg_elapsed is None or successes is None:
        return None
    try:
        return float(avg_elapsed) * int(successes)
    except (TypeError, ValueError):
        return None


def transfer_cell_text(row: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> str:
    lines = [
        provider_label(row, provider_labels),
        f"Median {fmt_mbps(row['median_mbps'])}",
        f"Avg {fmt_mbps(row['avg_mbps'])}",
        f"Total Time: {fmt_seconds(row.get('total_elapsed_seconds'))}",
    ]
    if row.get("median_elapsed_seconds") is not None:
        lines.append(f"Median time {fmt_seconds(row.get('median_elapsed_seconds'))}")
    return "\n".join(lines)


def traceroute_command(record: dict[str, Any]) -> str:
    command = str(record.get("trace_command") or "").strip()
    if command:
        return command
    endpoint = str(record.get("endpoint") or "endpoint")
    return f"traceroute -T -p 443 -n -w 3 -q 3 -m 30 {endpoint}"


def full_traceroute_block(record: dict[str, Any]) -> str:
    trace_output = str(record.get("trace_output") or record.get("last_hop") or "No raw traceroute output was captured in this JSONL record.")
    return f"$ {traceroute_command(record)}\n{trace_output}"


def transfer_rows(records: list[dict[str, Any]], record_type: str, file_set: str | None = None) -> list[dict[str, Any]]:
    rows = []
    for record in records:
        if record.get("record_type") != record_type:
            continue
        if file_set and record.get("file_set") != file_set:
            continue
        provider = provider_key(str(record.get("provider", "")))
        if provider not in PROVIDERS:
            continue
        rows.append({
            "provider": provider,
            "bucket": record.get("bucket", ""),
            "target_section": record.get("target_section", ""),
            "region": record.get("region", ""),
            "endpoint_url": record.get("endpoint_url", ""),
            "size_mib": float(record["file_size_mib"]),
            "attempts": record.get("attempt_count"),
            "successes": record.get("success_count"),
            "median_mbps": record.get("median_throughput_mbps"),
            "avg_mbps": record.get("avg_throughput_mbps"),
            "total_elapsed_seconds": summary_total_elapsed(record),
            "median_elapsed_seconds": record.get("median_elapsed_seconds"),
            "avg_elapsed_seconds": record.get("avg_elapsed_seconds"),
        })
    return rows


def latest_run_records(records: list[dict[str, Any]], record_type: str, file_set: str | None = None) -> list[dict[str, Any]]:
    candidates = [
        record for record in records
        if record.get("record_type") == record_type and (file_set is None or record.get("file_set") == file_set)
    ]
    run_ids = sorted({str(record.get("run_id", "")) for record in candidates if record.get("run_id")})
    if not run_ids:
        return candidates
    latest_run_id = run_ids[-1]
    return [record for record in candidates if str(record.get("run_id", "")) == latest_run_id]


def ranking_total_elapsed(row: dict[str, Any]) -> float:
    value = row.get("total_elapsed_seconds")
    if value is None or value == "":
        return float("inf")
    try:
        return float(value)
    except (TypeError, ValueError):
        return float("inf")


def ranked_by_size(rows: list[dict[str, Any]]) -> list[tuple[float, list[dict[str, Any]]]]:
    grouped = []
    for size in sorted({row["size_mib"] for row in rows}):
        members = [row for row in rows if row["size_mib"] == size]
        members.sort(key=lambda row: (
            ranking_total_elapsed(row),
            -(row.get("median_mbps") or 0),
            compact_provider(row["provider"]),
        ))
        grouped.append((size, members))
    return grouped


def file_size_count(members: list[dict[str, Any]]) -> int | None:
    counts = []
    for row in members:
        value = row.get("attempts")
        if value is None:
            value = row.get("successes")
        if value is None:
            continue
        try:
            counts.append(int(value))
        except (TypeError, ValueError):
            continue
    return max(counts) if counts else None


def file_size_label(size: float, members: list[dict[str, Any]]) -> str:
    count = file_size_count(members)
    return f"{fmt_size(size)} * {count}" if count else fmt_size(size)


def file_size_summary_labels(report_data: dict[str, Any]) -> list[str]:
    rows = report_data["upload_standard"] + report_data["upload_large"] + report_data["download"]
    labels = []
    for size in sorted({row["size_mib"] for row in rows}):
        labels.append(file_size_label(size, [row for row in rows if row["size_mib"] == size]))
    return labels


def total_time_chart_rows(rows: list[tuple[float, list[dict[str, Any]]]], provider_labels: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    totals: dict[str, dict[str, Any]] = {}
    for _size, members in rows:
        for row in members:
            total = ranking_total_elapsed(row)
            if total == float("inf"):
                continue
            key = provider_key(str(row.get("provider", ""))) or provider_chart_label(row, provider_labels)
            if key not in totals:
                totals[key] = {
                    "provider": row.get("provider", ""),
                    "label": provider_chart_label(row, provider_labels),
                    "fill": provider_bar_fill(row, provider_labels),
                    "total_seconds": 0.0,
                }
            totals[key]["total_seconds"] += total
    return sorted(totals.values(), key=lambda row: (row["total_seconds"], row["label"]))


def pct_delta(current: float, reference: float) -> float | None:
    if reference <= 0 or current == float("inf") or reference == float("inf"):
        return None
    return ((current - reference) / reference) * 100


def fmt_pct(value: float | None) -> str:
    if value is None:
        return "n/a"
    return f"{abs(value):.1f}%"


def filone_section_takeaway(title: str, rows: list[tuple[float, list[dict[str, Any]]]], provider_labels: dict[str, dict[str, Any]]) -> tuple[str, list[str]]:
    details = []
    for size, members in rows:
        if not members:
            continue
        filone_index = next((idx for idx, row in enumerate(members) if is_filone_row(row, provider_labels)), None)
        if filone_index is None:
            continue
        filone = members[filone_index]
        filone_time = ranking_total_elapsed(filone)
        rank = filone_index + 1
        if rank == 1:
            comparison = "ranked #1"
            if len(members) > 1:
                second_time = ranking_total_elapsed(members[1])
                delta = pct_delta(filone_time, second_time)
                comparison = f"ranked #1, {fmt_pct(delta)} faster than #2" if delta is not None else "ranked #1"
        else:
            fastest_time = ranking_total_elapsed(members[0])
            delta = pct_delta(filone_time, fastest_time)
            comparison = f"ranked #{rank}, {fmt_pct(delta)} slower than #1" if delta is not None else f"ranked #{rank}"
        details.append(f"{fmt_size(size)}: Fil.one {comparison}.")
    if not details:
        return title, ["Fil.one was not present in the loaded results."]
    return title, details


def filone_takeaway_sections(report_data: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> list[tuple[str, list[str]]]:
    sections = [
        ("Upload standard", ranked_by_size(report_data["upload_standard"])),
        ("Upload large", ranked_by_size(report_data["upload_large"])),
        ("Download", ranked_by_size(report_data["download"])),
    ]
    takeaways = [filone_section_takeaway(title, rows, provider_labels) for title, rows in sections]
    takeaways.append((
        "Methodology notes",
        [
            "Rankings are based on total elapsed time for each file-size group; lower total time is better.",
            "The Ookla baseline helps separate provider and route behavior from raw VM network capacity.",
            "All source data remains in JSONL so it can be re-used for later analysis or report generation.",
        ],
    ))
    return takeaways


def set_cell_text(cell, text: str, bold_first_line: bool = False, font_size: float = 9.0) -> None:
    cell.text = ""
    lines = text.split("\n")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.bold = (bold_first_line and idx == 0) or line.startswith("Total Time:")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths: list[float], header_fill: str = "F1F3F4") -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = False


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Normal"
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_monospace_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, "F8F9FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.splitlines() or [""]):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(7.0)
    style_table(table, [6.5], header_fill="F8F9FA")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = False


def add_ranked_table(doc: Document, title: str, rows: list[tuple[float, list[dict[str, Any]]]], provider_labels: dict[str, dict[str, Any]]) -> None:
    add_heading(doc, title, 2)
    if not rows:
        add_body(doc, "No matching summary records were found for this section.")
        return
    table = doc.add_table(rows=1, cols=5)
    for idx, header in enumerate(["File size / count", "Fastest", "2nd", "3rd", "4th"]):
        set_cell_text(table.rows[0].cells[idx], header, bold_first_line=True, font_size=8.5)
    for size, members in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], file_size_label(size, members), bold_first_line=True, font_size=8.5)
        for idx in range(4):
            text = "n/a"
            if idx < len(members):
                m = members[idx]
                text = transfer_cell_text(m, provider_labels)
            set_cell_text(cells[idx + 1], text, bold_first_line=True, font_size=7.2)
            if idx < len(members) and is_filone_row(members[idx], provider_labels):
                shade_cell(cells[idx + 1], FILONE_BRAND_FILL)
    style_table(table, [0.75, 1.43, 1.43, 1.43, 1.43])


def add_total_time_bar_chart(doc: Document, title: str, rows: list[tuple[float, list[dict[str, Any]]]], provider_labels: dict[str, dict[str, Any]]) -> None:
    chart_rows = total_time_chart_rows(rows, provider_labels)
    if not chart_rows:
        return
    add_body(doc, f"{title}: Total Time Ranking")
    max_total = max(row["total_seconds"] for row in chart_rows)
    table = doc.add_table(rows=1, cols=3 + BAR_SEGMENTS)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(["Rank", "Provider", "Total Time"]):
        set_cell_text(header_cells[idx], header, bold_first_line=True, font_size=7.4)
    for idx in range(BAR_SEGMENTS):
        set_cell_text(header_cells[3 + idx], "", font_size=6.0)
    for rank, row in enumerate(chart_rows, start=1):
        cells = table.add_row().cells
        set_cell_text(cells[0], str(rank), font_size=7.0)
        set_cell_text(cells[1], row["label"], bold_first_line=True, font_size=6.8)
        set_cell_text(cells[2], fmt_seconds(row["total_seconds"]), bold_first_line=True, font_size=7.0)
        active = max(1, round((row["total_seconds"] / max_total) * BAR_SEGMENTS)) if max_total else 0
        for idx in range(BAR_SEGMENTS):
            set_cell_text(cells[3 + idx], "", font_size=6.0)
            shade_cell(cells[3 + idx], row["fill"] if idx < active else "FFFFFF")
    style_table(table, [0.35, 1.75, 0.8] + [0.14] * BAR_SEGMENTS)


def add_network_table(doc: Document, records: list[dict[str, Any]]) -> None:
    add_heading(doc, "Node Network Baseline (Speedtest by Ookla)", 2)
    rows = [r for r in records if r.get("record_type") == "network_speedtest_summary"]
    if not rows:
        add_body(doc, "No Ookla network summary records were found.")
        return
    table = doc.add_table(rows=1, cols=5)
    for idx, header in enumerate(["Target", "Median ping", "Median download", "Median upload", "Packet loss"]):
        set_cell_text(table.rows[0].cells[idx], header, bold_first_line=True, font_size=8.5)
    for r in sorted(rows, key=lambda x: -(x.get("median_download_mbps") or 0)):
        cells = table.add_row().cells
        set_cell_text(cells[0], f"{r.get('target_city') or r.get('target_label')}\n{r.get('target_country') or ''}", bold_first_line=True, font_size=8)
        set_cell_text(cells[1], f"{r.get('median_ping_ms', 0):.2f} ms", font_size=8)
        set_cell_text(cells[2], fmt_mbps(r.get("median_download_mbps")), font_size=8)
        set_cell_text(cells[3], fmt_mbps(r.get("median_upload_mbps")), font_size=8)
        set_cell_text(cells[4], f"{r.get('max_packet_loss_percent', 0)}%", font_size=8)
    style_table(table, [1.25, 1.15, 1.35, 1.35, 1.0])


def add_traceroute_table(doc: Document, records: list[dict[str, Any]], endpoints: set[str]) -> None:
    add_heading(doc, "Provider Traceroutes", 2)
    rows = filter_traceroute_records(records, endpoints)
    if not rows:
        add_body(doc, "No provider traceroute records were found.")
        return

    rows.sort(key=lambda r: (r.get("provider") or "", r.get("region_location") or ""))
    table = doc.add_table(rows=1, cols=7)
    headers = ["Provider", "Region / location", "Endpoint", "Status", "Hops", "Total ms", "Resolved IPv4"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold_first_line=True, font_size=7.2)

    for r in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], str(r.get("provider") or "n/a"), bold_first_line=True, font_size=7.3)
        set_cell_text(cells[1], str(r.get("region_location") or "n/a"), font_size=6.8)
        set_cell_text(cells[2], str(r.get("endpoint") or "n/a"), font_size=6.2)
        set_cell_text(cells[3], str(r.get("status") or "n/a"), font_size=7.0)
        set_cell_text(cells[4], str(r.get("hop_count") if r.get("hop_count") is not None else "n/a"), font_size=7.0)
        set_cell_text(cells[5], fmt_ms(traceroute_total_ms(r)), font_size=7.0)
        set_cell_text(cells[6], fmt_trace_ips(r.get("resolved_ipv4_addresses")), font_size=6.2)

    style_table(table, [0.85, 1.05, 1.45, 0.55, 0.45, 0.75, 1.4])

    run_ids = sorted({str(r.get("run_id")) for r in rows if r.get("run_id")})
    if run_ids:
        add_body(doc, f"Traceroute source: TCP port 443 traceroute JSONL, latest included run_id {run_ids[-1]}. Total ms is the average RTT from the last responding hop in the traceroute output.")

    for r in rows:
        total_ms = fmt_ms(traceroute_total_ms(r))
        label = f"{r.get('provider') or 'Provider'} - {r.get('endpoint') or 'endpoint'} - total {total_ms}"
        add_body(doc, label)
        add_monospace_block(doc, full_traceroute_block(r))


def add_specs_table(doc: Document, args: argparse.Namespace) -> None:
    table = doc.add_table(rows=6, cols=2)
    specs = [
        ("Source provider", args.source_provider),
        ("Source location", args.source_location),
        ("Hostname", args.node_hostname),
        ("Network", args.node_network),
        ("Compute", args.node_compute),
        ("Memory", args.node_memory),
    ]
    for row, (key, value) in zip(table.rows, specs):
        set_cell_text(row.cells[0], key, bold_first_line=True, font_size=9)
        set_cell_text(row.cells[1], value, font_size=9)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_table(table, [1.55, 4.85], header_fill="FFFFFF")
    for row in table.rows:
        shade_cell(row.cells[0], "F1F3F4")


def load_report_data(data_dir: Path) -> dict[str, Any]:
    network_records = load_jsonl(data_dir / "network_speedtest_ookla_summary.jsonl")
    upload_records = load_jsonl(data_dir / "s3_upload_speedtest_summary.jsonl")
    download_records = load_jsonl(data_dir / "s3_download_speedtest_summary.jsonl")
    traceroute_records = load_jsonl(data_dir / "s3_provider_traceroutes.jsonl")

    if not upload_records:
        upload_records = load_jsonl(latest(data_dir, "s3_upload_speedtest_summary_*.jsonl", "s3_upload_speedtest_summary.jsonl"))
    if not download_records:
        download_records = load_jsonl(latest(data_dir, "s3_download_speedtest_summary_*.jsonl", "s3_download_speedtest_summary.jsonl"))
    if not traceroute_records:
        traceroute_records = load_jsonl(latest(data_dir, "s3_provider_traceroutes_*.jsonl", "s3_provider_traceroutes.jsonl"))

    upload_standard_records = latest_run_records(upload_records, "s3_upload_summary", "standard")
    upload_large_records = latest_run_records(upload_records, "s3_upload_summary", "large")
    download_latest_records = latest_run_records(download_records, "s3_download_summary")

    upload_standard = transfer_rows(upload_standard_records, "s3_upload_summary", "standard")
    upload_large = transfer_rows(upload_large_records, "s3_upload_summary", "large")
    if not upload_standard:
        upload_standard = [r for r in transfer_rows(latest_run_records(upload_records, "s3_upload_summary"), "s3_upload_summary") if r["size_mib"] <= 1024]
    if not upload_large:
        upload_large = [r for r in transfer_rows(latest_run_records(upload_records, "s3_upload_summary"), "s3_upload_summary") if r["size_mib"] >= 1024]
    download = transfer_rows(download_latest_records, "s3_download_summary")
    sizes = sorted({row["size_mib"] for row in upload_standard + upload_large + download})

    return {
        "network_records": network_records,
        "upload_standard": upload_standard,
        "upload_large": upload_large,
        "download": download,
        "traceroute_records": traceroute_records,
        "sizes": sizes,
    }


def provider_labels_text(provider_labels: dict[str, dict[str, Any]]) -> str:
    labels_by_name_region: dict[tuple[str, str, str], dict[str, Any]] = {}
    for label in provider_labels.values():
        name = str(label.get("name") or "").strip()
        region = str(label.get("region") or "").strip()
        location = str(label.get("location") or "").strip()
        if name:
            labels_by_name_region[(name, region, location)] = label
    if not labels_by_name_region:
        return "Provider titles use the provider, region, and bucket metadata from the loaded benchmark output."

    parts = []
    for name, region, location in sorted(labels_by_name_region):
        if region and location:
            parts.append(f"{name} ({region} | {location})")
        elif region:
            parts.append(f"{name} ({region})")
        else:
            parts.append(name)
    return "Provider titles from target config: " + ", ".join(parts) + "."


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(8)
    for style_name, size, before, after in [("Heading 1", 20, 20, 6), ("Heading 2", 16, 18, 6)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def create_doc(args: argparse.Namespace) -> Path:
    import_docx()
    data_dir = Path(args.data_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    report_data = load_report_data(data_dir)
    provider_labels = load_provider_labels(Path(args.targets))

    stamp = dt.datetime.now(dt.UTC).strftime("%Y%m%dT%H%M%SZ")
    docx = output_dir / f"s3_provider_speedtest_summary_{stamp}.docx"

    doc = Document()
    style_document(doc)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("S3 Provider Transfer Benchmark Summary")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(26)
    title_run.font.bold = False

    add_body(doc, f"Prepared from JSONL benchmark output in {data_dir}. Rankings in the transfer tables are ordered fastest to slowest from left to right using total elapsed time.")
    add_heading(doc, "Test Node", 2)
    add_specs_table(doc, args)
    add_network_table(doc, report_data["network_records"])
    add_heading(doc, "File Size Tests", 2)
    size_labels = ", ".join(file_size_summary_labels(report_data)) or "none"
    add_body(doc, "The benchmark file sizes represented in the loaded results are: " + size_labels + ".")

    upload_standard_ranked = ranked_by_size(report_data["upload_standard"])
    upload_large_ranked = ranked_by_size(report_data["upload_large"])
    download_ranked = ranked_by_size(report_data["download"])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_ranked_table(doc, "Upload Results: Small / Standard File Set", upload_standard_ranked, provider_labels)
    add_total_time_bar_chart(doc, "Upload Results: Small / Standard File Set", upload_standard_ranked, provider_labels)
    add_ranked_table(doc, "Upload Results: Large File Set", upload_large_ranked, provider_labels)
    add_total_time_bar_chart(doc, "Upload Results: Large File Set", upload_large_ranked, provider_labels)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_ranked_table(doc, "Download Results: All File Sizes", download_ranked, provider_labels)
    add_total_time_bar_chart(doc, "Download Results: All File Sizes", download_ranked, provider_labels)
    add_traceroute_table(doc, report_data["traceroute_records"], enabled_traceroute_endpoints(provider_labels))

    add_heading(doc, "Key Takeaways", 2)
    for section_title, details in filone_takeaway_sections(report_data, provider_labels):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(section_title)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
        for detail in details:
            sub = doc.add_paragraph(style="List Bullet 2")
            sub.paragraph_format.space_after = Pt(2)
            sub.paragraph_format.left_indent = Inches(0.35)
            sub_run = sub.add_run(detail)
            sub_run.font.name = "Arial"
            sub_run.font.size = Pt(10)

    doc.save(docx)
    return docx


def pdf_styles() -> dict[str, Any]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("ReportTitle", parent=base["Title"], fontName="Helvetica", fontSize=20, leading=24, alignment=0, spaceAfter=8),
        "heading": ParagraphStyle("ReportHeading", parent=base["Heading2"], fontName="Helvetica", fontSize=14, leading=17, spaceBefore=12, spaceAfter=6),
        "body": ParagraphStyle("ReportBody", parent=base["BodyText"], fontName="Helvetica", fontSize=9, leading=12, spaceAfter=6),
        "small": ParagraphStyle("ReportSmall", parent=base["BodyText"], fontName="Helvetica", fontSize=7.2, leading=8.5),
        "tiny": ParagraphStyle("ReportTiny", parent=base["BodyText"], fontName="Helvetica", fontSize=6.2, leading=7.2),
        "header": ParagraphStyle("ReportHeader", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=7.3, leading=8.5, alignment=1),
        "mono": ParagraphStyle("ReportMono", parent=base["Code"], fontName="Courier", fontSize=5.8, leading=6.7),
    }


def pdf_paragraph(text: Any, style: Any) -> Any:
    escaped_lines = []
    for line in str(text).split("\n"):
        escaped = html.escape(line)
        if line.startswith("Total Time:"):
            escaped = f"<b>{escaped}</b>"
        escaped_lines.append(escaped)
    return Paragraph("<br/>".join(escaped_lines), style)


def pdf_table(rows: list[list[Any]], widths: list[float], styles: dict[str, Any], repeat_header: bool = True, backgrounds: list[tuple[int, int, str]] | None = None) -> Any:
    converted = []
    for row_idx, row in enumerate(rows):
        row_style = styles["header"] if row_idx == 0 and repeat_header else styles["small"]
        converted.append([pdf_paragraph(cell, row_style) for cell in row])
    table = PdfTable(converted, colWidths=[width * inch for width in widths], repeatRows=1 if repeat_header else 0)
    table_style = [
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#DADCE0")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F3F4")) if repeat_header else ("BACKGROUND", (0, 0), (-1, -1), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    for row_idx, col_idx, fill in backgrounds or []:
        table_style.append(("BACKGROUND", (col_idx, row_idx), (col_idx, row_idx), colors.HexColor(f"#{fill}")))
    table.setStyle(TableStyle(table_style))
    return table


def pdf_add_heading(story: list[Any], text: str, styles: dict[str, Any]) -> None:
    story.append(Paragraph(html.escape(text), styles["heading"]))


def pdf_add_body(story: list[Any], text: str, styles: dict[str, Any]) -> None:
    story.append(Paragraph(html.escape(text), styles["body"]))


def pdf_add_specs_table(story: list[Any], args: argparse.Namespace, styles: dict[str, Any]) -> None:
    rows = [
        ["Source provider", args.source_provider],
        ["Source location", args.source_location],
        ["Hostname", args.node_hostname],
        ["Network", args.node_network],
        ["Compute", args.node_compute],
        ["Memory", args.node_memory],
    ]
    story.append(pdf_table(rows, [1.7, 7.7], styles, repeat_header=False))
    story.append(Spacer(1, 0.08 * inch))


def pdf_add_ranked_table(story: list[Any], title: str, rows: list[tuple[float, list[dict[str, Any]]]], styles: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> None:
    pdf_add_heading(story, title, styles)
    if not rows:
        pdf_add_body(story, "No matching summary records were found for this section.", styles)
        return
    table_rows = [["File size / count", "Fastest", "2nd", "3rd", "4th"]]
    backgrounds = []
    for size, members in rows:
        row = [file_size_label(size, members)]
        table_row_idx = len(table_rows)
        for idx in range(4):
            text = "n/a"
            if idx < len(members):
                member = members[idx]
                text = transfer_cell_text(member, provider_labels)
                if is_filone_row(member, provider_labels):
                    backgrounds.append((table_row_idx, idx + 1, FILONE_BRAND_FILL))
            row.append(text)
        table_rows.append(row)
    story.append(pdf_table(table_rows, [1.05, 2.05, 2.05, 2.05, 2.05], styles, backgrounds=backgrounds))


def pdf_add_total_time_bar_chart(story: list[Any], title: str, rows: list[tuple[float, list[dict[str, Any]]]], styles: dict[str, Any], provider_labels: dict[str, dict[str, Any]]) -> None:
    chart_rows = total_time_chart_rows(rows, provider_labels)
    if not chart_rows:
        return
    story.append(Paragraph(html.escape(f"{title}: Total Time Ranking"), styles["body"]))
    max_total = max(row["total_seconds"] for row in chart_rows)
    table_rows = [["Rank", "Provider", "Total Time"] + [""] * BAR_SEGMENTS]
    backgrounds = []
    for rank, row in enumerate(chart_rows, start=1):
        table_row_idx = len(table_rows)
        active = max(1, round((row["total_seconds"] / max_total) * BAR_SEGMENTS)) if max_total else 0
        table_rows.append([rank, row["label"], fmt_seconds(row["total_seconds"])] + [""] * BAR_SEGMENTS)
        for idx in range(active):
            backgrounds.append((table_row_idx, 3 + idx, row["fill"]))
    story.append(pdf_table(table_rows, [0.35, 1.8, 0.75] + [0.105] * BAR_SEGMENTS, styles, backgrounds=backgrounds))


def pdf_add_network_table(story: list[Any], records: list[dict[str, Any]], styles: dict[str, Any]) -> None:
    pdf_add_heading(story, "Node Network Baseline (Speedtest by Ookla)", styles)
    rows = [r for r in records if r.get("record_type") == "network_speedtest_summary"]
    if not rows:
        pdf_add_body(story, "No Ookla network summary records were found.", styles)
        return
    table_rows = [["Target", "Median ping", "Median download", "Median upload", "Packet loss"]]
    for row in sorted(rows, key=lambda x: -(x.get("median_download_mbps") or 0)):
        table_rows.append([
            f"{row.get('target_city') or row.get('target_label')}\n{row.get('target_country') or ''}",
            f"{row.get('median_ping_ms', 0):.2f} ms",
            fmt_mbps(row.get("median_download_mbps")),
            fmt_mbps(row.get("median_upload_mbps")),
            f"{row.get('max_packet_loss_percent', 0)}%",
        ])
    story.append(pdf_table(table_rows, [2.0, 1.3, 1.7, 1.7, 1.1], styles))


def pdf_trace_output(text: str) -> str:
    return "\n".join(text.splitlines() or [""])


def pdf_add_traceroutes(story: list[Any], records: list[dict[str, Any]], styles: dict[str, Any], endpoints: set[str]) -> None:
    pdf_add_heading(story, "Provider Traceroutes", styles)
    rows = filter_traceroute_records(records, endpoints)
    if not rows:
        pdf_add_body(story, "No provider traceroute records were found.", styles)
        return
    rows.sort(key=lambda r: (r.get("provider") or "", r.get("region_location") or ""))
    table_rows = [["Provider", "Region / location", "Endpoint", "Status", "Hops", "Total ms", "Resolved IPv4"]]
    for row in rows:
        table_rows.append([
            row.get("provider") or "n/a",
            row.get("region_location") or "n/a",
            row.get("endpoint") or "n/a",
            row.get("status") or "n/a",
            row.get("hop_count") if row.get("hop_count") is not None else "n/a",
            fmt_ms(traceroute_total_ms(row)),
            fmt_trace_ips(row.get("resolved_ipv4_addresses")),
        ])
    story.append(pdf_table(table_rows, [0.9, 1.2, 2.2, 0.75, 0.55, 0.8, 1.6], styles))

    run_ids = sorted({str(r.get("run_id")) for r in rows if r.get("run_id")})
    if run_ids:
        pdf_add_body(story, f"Traceroute source: TCP port 443 traceroute JSONL, latest included run_id {run_ids[-1]}. Total ms is the average RTT from the last responding hop in the traceroute output.", styles)

    for row in rows:
        total_ms = fmt_ms(traceroute_total_ms(row))
        label = f"{row.get('provider') or 'Provider'} - {row.get('endpoint') or 'endpoint'} - total {total_ms}"
        story.append(Paragraph(html.escape(label), styles["body"]))
        story.append(Preformatted(pdf_trace_output(full_traceroute_block(row)), styles["mono"]))
        story.append(Spacer(1, 0.06 * inch))


def create_pdf(args: argparse.Namespace) -> Path:
    import_reportlab()
    data_dir = Path(args.data_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    report_data = load_report_data(data_dir)
    provider_labels = load_provider_labels(Path(args.targets))

    stamp = dt.datetime.now(dt.UTC).strftime("%Y%m%dT%H%M%SZ")
    pdf = output_dir / f"s3_provider_speedtest_summary_{stamp}.pdf"
    doc = SimpleDocTemplate(
        str(pdf),
        pagesize=landscape(letter),
        leftMargin=0.45 * inch,
        rightMargin=0.45 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
    )
    styles = pdf_styles()
    story: list[Any] = []

    story.append(Paragraph("S3 Provider Transfer Benchmark Summary", styles["title"]))
    pdf_add_body(story, f"Prepared from JSONL benchmark output in {data_dir}. Rankings in the transfer tables are ordered fastest to slowest from left to right using total elapsed time.", styles)
    pdf_add_heading(story, "Test Node", styles)
    pdf_add_specs_table(story, args, styles)
    pdf_add_network_table(story, report_data["network_records"], styles)
    pdf_add_heading(story, "File Size Tests", styles)
    sizes = ", ".join(file_size_summary_labels(report_data)) or "none"
    pdf_add_body(story, "The benchmark file sizes represented in the loaded results are: " + sizes + ".", styles)

    upload_standard_ranked = ranked_by_size(report_data["upload_standard"])
    upload_large_ranked = ranked_by_size(report_data["upload_large"])
    download_ranked = ranked_by_size(report_data["download"])

    story.append(PageBreak())
    pdf_add_ranked_table(story, "Upload Results: Small / Standard File Set", upload_standard_ranked, styles, provider_labels)
    pdf_add_total_time_bar_chart(story, "Upload Results: Small / Standard File Set", upload_standard_ranked, styles, provider_labels)
    pdf_add_ranked_table(story, "Upload Results: Large File Set", upload_large_ranked, styles, provider_labels)
    pdf_add_total_time_bar_chart(story, "Upload Results: Large File Set", upload_large_ranked, styles, provider_labels)

    story.append(PageBreak())
    pdf_add_ranked_table(story, "Download Results: All File Sizes", download_ranked, styles, provider_labels)
    pdf_add_total_time_bar_chart(story, "Download Results: All File Sizes", download_ranked, styles, provider_labels)
    pdf_add_traceroutes(story, report_data["traceroute_records"], styles, enabled_traceroute_endpoints(provider_labels))

    pdf_add_heading(story, "Key Takeaways", styles)
    for section_title, details in filone_takeaway_sections(report_data, provider_labels):
        story.append(Paragraph("&bull; <b>" + html.escape(section_title) + "</b>", styles["body"]))
        for detail in details:
            story.append(Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;&bull; " + html.escape(detail), styles["body"]))

    doc.build(story)
    return pdf


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a DOCX or PDF report from benchmark JSONL output.")
    parser.add_argument("--data-dir", default="/dataoutput", help="Directory containing JSONL benchmark output")
    parser.add_argument("--output-dir", default="/dataoutput/reports", help="Directory for generated reports")
    parser.add_argument("--targets", default="/testfiles/s3_targets.ini", help="INI file with bucket/provider targets used for provider labels")
    parser.add_argument("--format", choices=["docx", "pdf", "both"], default="docx", help="Report output format")
    parser.add_argument("--source-provider", "--node-name", dest="source_provider", default="", help="Source node provider/name; prompted when omitted in an interactive terminal")
    parser.add_argument("--source-location", "--node-location", dest="source_location", default="", help="Source node location; prompted when omitted in an interactive terminal")
    parser.add_argument("--node-hostname", default="", help="Source node hostname; auto-detected with hostname when omitted")
    parser.add_argument("--node-network", default=os.environ.get("SOURCE_NODE_NETWORK", "unknown"), help="Source node network description")
    parser.add_argument("--node-compute", default="", help="Source node compute description; auto-detected with nproc when omitted")
    parser.add_argument("--node-memory", default="", help="Source node memory description; auto-detected with free/procfs when omitted")
    parser.add_argument("--no-prompt", action="store_true", help="Do not prompt for source provider/location; use flags, env vars, or fallback values")
    args = parser.parse_args()
    args = resolve_source_context(args)
    outputs = []
    if args.format in {"docx", "both"}:
        outputs.append(create_doc(args))
    if args.format in {"pdf", "both"}:
        outputs.append(create_pdf(args))
    for output in outputs:
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
